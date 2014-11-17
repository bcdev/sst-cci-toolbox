__author__ = 'Ralf Quast'

from docx import Document


class Car:
    def __init__(self):
        self.document = Document()
        self.add_static_contents()

    def add_static_contents(self):
        with open('title.txt', 'r') as f:
            title = f.read()
        self.document.add_heading(title, 0)
        with open('contents.txt', 'r') as f:
            contents = f.readlines()
        for line in contents:
            resource_name = line.replace('\n', '')
            """:type : str"""
            if resource_name.startswith('h'):
                level = resource_name.count('h')
                with open(resource_name, 'r') as f:
                    heading = f.read()
                self.document.add_heading(heading, level)
            elif resource_name.startswith('p'):
                with open(resource_name, 'r') as f:
                    paragraph = f.read()
                self.document.add_paragraph(paragraph)
            elif resource_name.startswith('f'):
                self.document.add_picture(resource_name)
            else:
                self.document.add_page_break()

    def get_document(self):
        """

        :rtype : Document
        """
        return self.document